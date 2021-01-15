from docx import Document
# import json
import pymysql.cursors
# from pathlib import Path
from django.conf import settings
from .models import Upload

res = dict()

filepath = settings.MEDIA_ROOT


def savetodb():
    print("Process Started")
    document = Document(filepath + '/data.docx')
    tables = document.tables

    tab = []

    for u in tables:
        h = [[row.cells[a].paragraphs[0].text for row in u.columns]
             for a in range(len(u.rows))]
        dic = {x: y for (x, y) in zip(h[0], zip(*h[1:]))}
        keys = ['Receipt No.', 'Completion Time',
                'Details', 'Paid In', 'Withdrawn', 'Balance']
        tab.append({key: dic[key] for key in dic if key in keys})

    for dict in tab:
        for list in dict:
            if list in res:
                res[list] += (dict[list])
            else:
                res[list] = dict[list]

    # with open('test.json', 'w') as fd:
        # json.dump(res, fd)

    # sudo /opt/lampp/lampp stop

    # with open('test.json', 'r') as file:
        # data = json.load(file)

    data = res

    # Connect to the database.
    connection = pymysql.connect(host='localhost',
                                 user='root',
                                 password='mat@1234',
                                 db='japhe',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)

    print ("connect successful!!")

    with connection:
        with connection.cursor() as cursor:
            # Create a new record
            sql = "INSERT INTO `transactions` (`Receipt_No`, `Completion_Time`, `Details`, `Paid_In`, `Withdrawn`, `Balance`) VALUES (%s, %s, %s, %s, %s, %s)"
            for a in range(len(data['Receipt No.'])):

                try:
                    cursor.execute(
                        sql, (data['Receipt No.'][a], data['Completion Time'][a], data['Details'][a], data['Paid In'][a], data['Withdrawn'][a], data['Balance'][a]))

                except:
                    print('DB upto date fox')
                    break

        # connection is not autocommit by default. So you must commit to save
        # your changes.
        connection.commit()


        with connection.cursor() as cursor:
            # Read a single record
            sql = "SELECT `Receipt_No` FROM `transactions` WHERE `Paid_In`=%s"
            cursor.execute(sql, ('500.00',))
            result = cursor.fetchone()
            print(result)
            # Upload.objects.all().delete()


if __name__ == '__main__':
	savetodb()
